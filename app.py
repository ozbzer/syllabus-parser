import streamlit as st
from pypdf import PdfReader
from openai import OpenAI
import os
from streamlit.errors import StreamlitSecretNotFoundError
import json 
from datetime import datetime, timedelta

st.title("Syllabus Parser")
st.info("Turns your syllabus into full study plan + calendar automatically")

# Upload file
def upload_file():
    uploaded_file = st.file_uploader("Only accept PDF & DOCX Syllabus.")
    return uploaded_file

uploaded_file = upload_file()

if uploaded_file: # checks whether a file was uploaded
    st.text(uploaded_file.name)
else:
    st.text("Please upload your syllabus")

def generate_ai_today_text(syllabus_text):
    prompt = f"""
Read the syllabus text below and extract all important academic deadlines and events that should be added to a student's personal calendar.

Look for:
- assignments
- quizzes
- exams
- projects
- presentations
- labs
- essays
- reports
- readings only if they have a due date
- important class deadlines
- final exams or major assessments
- labs only when something must be submitted, completed, or prepared by a certain date
- waivers or required forms only when they have a due date
- graded attendance only if the syllabus clearly treats it as an assessed requirement with a date

Do NOT include:
- lecture topics
- attendance
- weekly themes
- class meetings
- field trips by themselves
- artist talks by themselves
- readings unless explicitly due
- general course activities
- office hours
- location information
- announcements without a deadline

Return only valid JSON.
Do not return explanations, markdown, headings, or extra text.

Return a JSON array.
Each item in the array must follow this format:
{{
  "title": "string",
  "type": "assignment | quiz | exam | project | presentation | lab | reading | reflection | paper | essay | report | lab | deadline",
  "date": "YYYY-MM-DD or null",
  "time": "HH:MM or null",
  "description": "short helpful detail or null",
  "source_section": "evaluation | schedule | other"
  "evidence_text": "exact short quote or excerpt from the syllabus supporting the date"
}}

Rules:
- Only include items that have a clear due date or scheduled date.
- Do not guess missing dates or times.
- If the date is missing, do not include the item.
- If the time is missing, use null.
- Keep titles short and student-friendly.
- Sort items by date from earliest to latest.
- Don't include items that are not calendar-worthy (e.g., "Office Hours", "Syllabus Overview")
- If a syllabus contains both a course evaluation section and a weekly schedule, prioritize dates from the evaluation/grading section for quizzes, exams, projects, presentations, reflections, and other graded work.
- Do not use lecture-topic dates as assignment dates unless the syllabus clearly says the item is due on that date.
- For each item, label where the date came from:
  - "evaluation" if it came from a grading/evaluation/assessment table or section
  - "schedule" if it came from the weekly course calendar or class schedule
  - "other" otherwise
- If the same graded item appears more than once with conflicting dates, prefer the one from "evaluation".
- Only include an item if you can provide a short exact supporting excerpt from the syllabus.
- If the date is uncertain or conflicting, prefer the item with clearer evidence from an evaluation/grading section.
- If no valid calendar items are found, return [].

Syllabus text:
{syllabus_text}
"""
    
    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )
    return response.output_text

def generate_prep_events(title, due_date, task_type):
    prep_events = [] # Create an empty list where I will store all the prep events
    clean_title = title.replace("Due", "").strip()
    clean_title = clean_title.replace("Research Draft", "Proposal")
    clean_title = clean_title.replace("Take-Home Exam", "Exam")
    # Step 1: decide prep steps 
    if task_type == "assignment":
        steps = [("Research", 7), ("Outline", 5), ("Draft", 3), ("Final Edit", 1)]
    elif task_type == "exam":
        steps = [("Study Session 1", 6), ("Study Session 2", 3), ("Review", 1)]
    elif task_type == "quiz":
        steps = [("Review Notes", 3), ("Practice Questions", 1)]
    elif task_type == "project":
        steps = [("Plan Project", 10), ("Work Session 1", 7), ("Work Session 2", 4), ("Final Touches", 1)]
    elif task_type == "presentation":
        steps = [("Research", 7), ("Slides Draft", 4), ("Practice", 2)]
    elif task_type == "reflection":
        steps = [("Review Notes", 3), ("Draft Reflection", 1)]
    elif task_type == "lab":
        steps = [("Prepare Materials", 2), ("Review Instructions", 1)]
    elif task_type == "deadline":
        steps = [("Start Task", 5), ("Finish Task", 1)]
    else:
        steps = []

    due = datetime.strptime(due_date, "%Y-%m-%d")

    for step_name, days_before in steps:
        prep_date = due - timedelta(days=days_before)

        event_title = f"{step_name} for {clean_title}"

        words = event_title.split()
        event_title = " ".join(dict.fromkeys(words))

        prep_events.append({
            "title": event_title,
            "date": prep_date.strftime("%Y-%m-%d")
        }) 
    return prep_events
   
analyze_button = st.button("Analyze Syllabus")

def prefer_evaluation_dates(calendar_items):
    final = {}
    for item in calendar_items:
        key = item.get("title", "").strip().lower()
        if not key:
            continue

        # If same item appears twice → prefer evaluation
        if key not in final or item.get("source_section") == "evaluation":
            final[key] = item

    return list(final.values())

if analyze_button:
    if uploaded_file:
        st.text("Analyzing syllabus...")

        syllabus_text = ""

        if uploaded_file.name.endswith(".pdf"):
            pdf_reader = PdfReader(uploaded_file)

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    syllabus_text += page_text + "\n"

        elif uploaded_file.name.endswith(".docx"):
            from docx import Document

            doc = Document(uploaded_file)

            for para in doc.paragraphs:
                if para.text.strip():
                    syllabus_text += para.text.strip() + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)

                    if row_text:
                        syllabus_text += " | ".join(row_text) + "\n"

        if syllabus_text.strip():
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                try:
                    api_key = st.secrets.get("OPENAI_API_KEY")
                except StreamlitSecretNotFoundError:
                    api_key = None

            if api_key:
                client = OpenAI(api_key=api_key)

                ai_response = generate_ai_today_text(syllabus_text)
                try:
                    calendar_items = json.loads(ai_response)
                    calendar_items = prefer_evaluation_dates(calendar_items) 
                except json.JSONDecodeError:
                    st.error("AI response was not valid JSON.")
                    st.text_area("Raw AI response", ai_response, height=300)
                    st.stop()
                st.session_state["calendar_items"] = calendar_items 
                # New logic comes here 
                prep_events_lists = []
                st.session_state["prep_events_lists"] = prep_events_lists
                for task in calendar_items:
                    title = task["title"]
                    due_date = task["date"]
                    task_type = task["type"]

                    if not due_date:
                        continue

                    prep_events = generate_prep_events(title, due_date, task_type)
                    prep_events_lists.extend(prep_events) # combined them into one list
                
                # add preview here
                st.subheader("Deadlines Found")
                for event in calendar_items:
                    title = event["title"]
                    title = title.replace(" Due", "").strip()
                    date = event["date"]
                    if not date:
                        continue 
                    formatted_date = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
                    st.write(f"**{title}** — {formatted_date}")

                st.subheader("Your Study Plan")
                for prep_event in prep_events_lists:
                    title = prep_event["title"]
                    date = prep_event["date"]
                    if not date:
                        continue
                    formatted_date = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
                    st.write(f"**{title}** — {formatted_date}")
                    
                calendar_text = "BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//Syllabus Parser//EN\n"
                for item in calendar_items:
                    title = item["title"]
                    date = item["date"]
                    if not date:
                        continue
                    description = item["description"]

                    event_text = f"""BEGIN:VEVENT
SUMMARY:{title}
DTSTART;VALUE=DATE:{date.replace("-", "")}
DESCRIPTION:{description if description else ""}
END:VEVENT
"""
                    calendar_text += event_text

                for prep_event in prep_events_lists:
                    title = prep_event["title"] 
                    date = prep_event["date"]
                    event_text = f"""BEGIN:VEVENT
SUMMARY:{title}
DTSTART;VALUE=DATE:{date.replace("-", "")}
END:VEVENT
"""
                    calendar_text += event_text
                
                calendar_text += "END:VCALENDAR\n"
                st.success("Calendar created successfully. Download it below.")
                st.download_button(
                    "Download Calendar",
                    data=calendar_text,
                    file_name="syllabus_calendar.ics",
                    key="download_calendar_button",
                )
                st.info("Tip: Review your calendar before using — AI may miss or misinterpret some details.")

            else:
                st.warning("OpenAI API key not found. Please set OPENAI_API_KEY.")
        else:
            st.warning("Could not extract text from this file.Try another syllabus or format.")
    else:
        st.warning("Please upload a syllabus first")


